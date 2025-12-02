from fastapi import APIRouter, Depends, HTTPException
from dependencies import get_current_user
from app.utils.database import supabase
from .user import get_user_info, get_student_type
from app.utils.openai_client import ask_openai, ask_gemini
from pydantic import BaseModel


from PyPDF2 import PdfReader
from docx import Document
from PIL import Image
import pytesseract
import io
import datetime
import json

router = APIRouter()


def extract_text_from_file(data: bytes, filename: str) -> str:
    fileName = filename.lower()

    if fileName.endswith(".pdf"):
        reader = PdfReader(io.BytesIO(data))
        pages = [page.extract_text() or "" for page in reader.pages]
        return "\n".join(pages)

    elif fileName.endswith(".docx"):
        doc = Document(io.BytesIO(data))
        texts = [p.text for p in doc.paragraphs]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texts.append(cell.text)
        return "\n".join(texts)

    elif (
        fileName.endswith(".png")
        or fileName.endswith(".jpeg")
        or fileName.endswith(".jpg")
        or fileName.endswith(".tiff")
    ):
        img = Image.open(io.BytesIO(data))
        # ensure RGB or grayscale:
        if img.mode not in ("RGB", "L"):
            img = img.convert("RGB")
        return pytesseract.image_to_string(img)

    else:
        raise HTTPException(
            status_code=400,
            detail="Unsupported file type; only PDF, DOCX, and JPEG/PNG/TIFF images are allowed.",
        )


class AnalyseReportRequest(BaseModel):
    file_path: str


@router.post("/analyse")
async def analyse_report(request: AnalyseReportRequest, user=Depends(get_current_user)):
    student_type = await get_student_type(user)
    user_info = await get_user_info(user, student_type)

    if student_type == "high_school":
        table = "student_school_data"
        report_table = "school_report_analysis"
    else:
        table = "student_uni_data"
        report_table = "transcript_analysis"

    # Use the path from the request instead of querying
    path = request.file_path

    if not path:
        raise HTTPException(status_code=400, detail="File path is required")

    download = supabase.storage.from_("reports").download(path)

    if not download:
        raise HTTPException(status_code=400, detail="Could not Download")

    raw = download
    report_text = extract_text_from_file(raw, path)

    if student_type == "high_school":
        prompt = f"""
          You are an expert high‐school academic analyst. Analyse the following student school report and return a detailed analysis on the report. 
          This report will later be used for openAI to read again so return in markdown:

          - **top_subjects** (array of strings): Subjects with the highest marks (2-3 subjects including the results/marks/rank if available).  
          - **bottom_subjects** (array of strings): Subjects with the lowest marks (1-2 subjects).  
          - **strengths** (array of strings): short phrases describing the student’s strongest areas.  
          - **weaknesses** (array of strings): short phrases describing areas needing improvement.  
          - **Evaluation** (string): A strong detailed explanation to what you analysed and can determine based on the results.
          - **Recommendation (string): Recommend Study Habits, which areas to focus on or prioritise and what to do heading forward.

          **High-School Report:**  
          {report_text}
          
          Return  **only** a single valid JSON object.
        """
    else:
        prompt = f"""
        You are a seasoned UNSW academic advisor. Analyse the following university transcript text and extract key academic information.

        Return the result as a single valid JSON object. Do not use markdown or explanations. Only include fields if the information is clearly present in the transcript.

        The JSON must include the following fields:

        - **full_name** (string): Student’s full name  
        - **student_id** (string): UNSW student ID  
        - **degree_program** (string): Full name of the degree program (e.g. "Bachelor of Engineering (Honours)")  
        - **major** (string): Name of the major or specialisation (e.g. "Computer Engineering")  
        - **wam** (number): Final Weighted Average Mark  
        - **uoc_attempted** (number): Total Units of Credit attempted  
        - **uoc_completed** (number): Total Units of Credit successfully completed  
        - **term_summaries** (array of objects): For each term:
        - `term` (string): e.g. "Term 1 2021"
        - `wam` (number): WAM for that term if present
        - `courses` (array): Each course with:
            - `code`: Course code (e.g. "COMP1511")
            - `title`: Course title
            - `uoc`: Units of Credit
            - `mark`: Mark (if available)
            - `grade`: Grade (e.g. "HD", "DN", "PS")

        Also include:

        - **high_achievements** (array of strings): Courses with high marks or HDs  
        - **low_performance** (array of strings): Courses with lowest grades or weak performance  
        - **strengths** (array of strings): Short phrases summarizing strong academic areas  
        - **weaknesses** (array of strings): Areas needing improvement  
        - **evaluation** (string): A concise paragraph summarizing overall academic standing  
        - **recommendation** (string): Specific advice on what to improve or prioritize going forward

        Return only the JSON result. Do not include markdown, explanation, or commentary.

        Transcript Text:
        {report_text}
        """

    ai_output_str = ask_gemini(prompt)
    text = ai_output_str.strip()
    if text.startswith("```"):
        # remove ```json or ``` at top/bottom
        text = text.strip("```json").strip("```").strip()

    # Parse it into a native dict
    try:
        ai_output = json.loads(text)
    except json.JSONDecodeError as e:
        raise HTTPException(500, f"Could not parse LLM output as JSON: {e}")

    # Upsert in one go (avoids having to check update vs insert yourself)
    upsert_payload = {
        "user_id": user.id,
        "analysis": ai_output,  # dict → JSONB
        "report_path": user_info["report_path"],
        "analysed_at": datetime.datetime.now().isoformat(),
    }

    resp = supabase.table(report_table).upsert(upsert_payload).execute()
    if not resp:
        raise HTTPException(500, f"DB upsert failed: {resp.error.message}")

    return {"analysis": ai_output}
